import os
import pandas as pd
from django.utils import timezone
from docx import Document
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
from .models import Candidate, Template, OfferLetter
from .google_service import GoogleDocsService

def process_docx_template(template_path, output_path, data):
    """Process DOCX template with placeholder replacement"""
    try:
        doc = Document(template_path)
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    full_text = paragraph.text
                    new_text = full_text.replace(key, str(value))
                    
                    # Clear all existing runs
                    for i in range(len(paragraph.runs) - 1, -1, -1):
                        paragraph._element.remove(paragraph.runs[i]._element)
                    
                    # Add new run with modified text
                    paragraph.add_run(new_text)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key in cell.text:
                            full_text = cell.text
                            new_text = full_text.replace(key, str(value))
                            
                            # Reset cell content
                            cell._element.xml = '<w:tc><w:p/></w:tc>'
                            cell.paragraphs[0].add_run(new_text)
        
        # Save the document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        return False

def convert_to_pdf(docx_path: str) -> str:
    """
    Convert DOCX file to PDF using LibreOffice in headless mode.
    
    Args:
        docx_path: Path to the DOCX file to convert
        
    Returns:
        str: Path to the generated PDF file
        
    Raises:
        subprocess.CalledProcessError: If LibreOffice conversion fails
        FileNotFoundError: If DOCX file or LibreOffice is not found
        Exception: For other conversion errors
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    # Get directory and filename without extension
    docx_dir = os.path.dirname(docx_path)
    docx_filename = os.path.basename(docx_path)
    pdf_filename = os.path.splitext(docx_filename)[0] + '.pdf'
    pdf_path = os.path.join(docx_dir, pdf_filename)
    
    try:
        # Run LibreOffice in headless mode for conversion
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', docx_dir,
            docx_path
        ]
        
        # Execute conversion
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60  # 60 second timeout for LibreOffice
        )
        
        if result.returncode != 0:
            error_msg = result.stderr.strip() if result.stderr else result.stdout.strip()
            raise subprocess.CalledProcessError(
                result.returncode, 
                cmd, 
                result.stdout, 
                result.stderr
            )
        
        # Check if PDF was created
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF conversion completed but file not found: {pdf_path}")
        
        return pdf_path
        
    except subprocess.TimeoutExpired:
        raise Exception(f"PDF conversion timed out after 60 seconds for file: {docx_path}")
    
    except subprocess.CalledProcessError as e:
        error_details = e.stderr.strip() if e.stderr else e.stdout.strip()
        raise Exception(f"LibreOffice conversion failed: {error_details}")
    
    except Exception as e:
        raise Exception(f"Error converting DOCX to PDF: {str(e)}")

def generate_offer_letter(candidate, template):
    """Generate offer letter for a candidate using Google Docs API"""
    try:
        # Initialize Google Docs service
        google_service = GoogleDocsService()
        
        # Prepare data for template
        data = {
            '{{name}}': candidate.name,
            '{{phone}}': candidate.phone,
            '{{email}}': candidate.email,
            '{{role}}': candidate.role,
            '{{letter_date}}': candidate.letter_date.strftime('%d %B %Y'),
            '{{joining_date}}': candidate.joining_date.strftime('%d %B %Y'),
            '{{work_id}}': candidate.work_id,
        }
        
        # Check if template has Google Doc ID
        if template.google_doc_id:
            # Use Google Docs API with smart method (no storage quota, proper replacement)
            pdf_file = google_service.generate_offer_pdf_smart(
                template.google_doc_id,
                data,
                candidate.name
            )
            
            # Save offer letter record with PDF
            offer_letter = OfferLetter.objects.create(
                candidate_id=candidate.id,
                template_id=template.id,
                candidate_work_id=candidate.work_id,
                template_name=template.name,
                pdf_file=pdf_file
            )
            
            # Update candidate status
            candidate.status = 'offer_generated'
            candidate.save()
            
            return offer_letter
        
        else:
            # Fallback to local DOCX processing (if Google Doc ID not set)
            return generate_offer_letter_fallback(candidate, template, data)
            
    except Exception as e:
        print(f"Error generating offer letter: {str(e)}")
        return None

def generate_offer_letter_fallback(candidate, template, data):
    """Fallback method using local DOCX processing"""
    try:
        # Generate temporary DOCX filename
        temp_docx_filename = f"temp_offer_{candidate.work_id}.docx"
        temp_docx_path = os.path.join(settings.MEDIA_ROOT, 'temp', temp_docx_filename)
        
        # Generate final PDF filename
        pdf_filename = f"offer_{candidate.work_id}.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, 'offer_letters', 'pdf', pdf_filename)
        
        # Ensure directories exist
        os.makedirs(os.path.dirname(temp_docx_path), exist_ok=True)
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        
        # Process template to temporary DOCX
        template_path = template.file.path
        success = process_docx_template(template_path, temp_docx_path, data)
        
        if success:
            try:
                # Convert DOCX to PDF using LibreOffice
                pdf_path = convert_to_pdf(temp_docx_path)
                
                if os.path.exists(pdf_path):
                    # Save offer letter record with PDF only
                    with open(pdf_path, 'rb') as pdf_file:
                        offer_letter = OfferLetter.objects.create(
                            candidate_id=candidate.id,
                            template_id=template.id,
                            candidate_work_id=candidate.work_id,
                            template_name=template.name,
                            pdf_file=ContentFile(pdf_file.read(), pdf_filename)
                        )
                    
                    # Clean up temporary files
                    if os.path.exists(temp_docx_path):
                        os.remove(temp_docx_path)
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                    
                    # Update candidate status
                    candidate.status = 'offer_generated'
                    candidate.save()
                    
                    return offer_letter
                else:
                    raise Exception(f"PDF file not created after conversion: {pdf_path}")
                    
            except Exception as conversion_error:
                # Clean up temporary DOCX file on conversion failure
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                raise Exception(f"PDF conversion failed for {candidate.work_id}: {str(conversion_error)}")
        
        return None
        
    except Exception as e:
        print(f"Error in fallback offer letter generation: {str(e)}")
        return None

def get_template_for_role(role):
    """Get appropriate template for a role (supports custom roles)"""
    # First try to find exact role match
    template = Template.objects.filter(role__iexact=role, is_active=True).first()
    if template:
        return template
    
    # Convert role display name to role key for backward compatibility
    role_mapping = {
        'Frontend Developer': 'frontend',
        'Backend Developer': 'backend',
        'Machine Learning Engineer': 'machine_learning',
        'Full Stack Developer': 'full_stack',
        'UI/UX Designer': 'ui_ux',
        'Digital Marketing': 'digital_marketing',
        'PR Specialist': 'pr',
        'Content Writer': 'content',
        'Video Editor': 'video_editor',
    }
    
    role_key = role_mapping.get(role, None)
    
    # Try legacy role key mapping
    if role_key:
        template = Template.objects.filter(role=role_key, is_active=True).first()
        if template:
            return template
    
    # If no template found, return None
    return None
    
    # Fallback to general template
    return Template.objects.filter(role__isnull=True, is_active=True).first()

def process_bulk_upload(file, created_by):
    """Process bulk upload of candidates from CSV/Excel"""
    try:
        # Read file based on extension
        if file.name.endswith('.csv'):
            # Read CSV with phone numbers as strings to preserve leading zeros
            df = pd.read_csv(file, dtype={'phone': str})
        elif file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file, dtype={'phone': str})
        else:
            return None, "Unsupported file format"
        
        # Validate required columns
        required_columns = ['name', 'email', 'phone', 'role', 'joining_date']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return None, f"Missing required columns: {', '.join(missing_columns)}"
        
        # Process candidates
        candidates = []
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Validate email format
                email = str(row['email']).strip()
                if '@' not in email or '.' not in email:
                    errors.append(f"Row {index + 1}: Invalid email format")
                    continue
                
                # Validate phone (basic check)
                phone = str(row['phone']).strip()
                # Remove any non-digit characters for validation
                phone_digits = ''.join(filter(str.isdigit, phone))
                if len(phone_digits) < 10:
                    errors.append(f"Row {index + 1}: Phone number too short (minimum 10 digits)")
                    continue
                
                # Validate role
                role = str(row['role']).strip()
                valid_roles = [choice[0] for choice in Candidate.ROLE_CHOICES] + \
                           [choice[1] for choice in Candidate.ROLE_CHOICES]
                if role not in valid_roles:
                    errors.append(f"Row {index + 1}: Invalid role '{role}'. Valid roles: {', '.join(valid_roles)}")
                    continue
                
                # Validate joining date
                try:
                    joining_date = pd.to_datetime(row['joining_date']).date()
                    current_date = timezone.now().date()
                    
                    if joining_date <= current_date:
                        errors.append(f"Row {index + 1}: Joining date must be in the future (current: {current_date}, provided: {joining_date})")
                        continue
                except Exception as date_error:
                    errors.append(f"Row {index + 1}: Invalid joining date format: {str(date_error)}")
                    continue
                
                candidate_data = {
                    'name': str(row['name']).strip(),
                    'email': email,
                    'phone': phone,
                    'role': role,
                    'letter_date': timezone.now().date(),
                    'joining_date': joining_date,
                    'created_by_id': created_by.id,
                    'created_by_username': created_by.username,
                }
                
                # Create candidate (work_id will be auto-generated)
                candidate = Candidate(**candidate_data)
                candidate.save()
                candidates.append(candidate)
                
            except Exception as e:
                errors.append(f"Row {index + 1}: {str(e)}")
        
        return candidates, errors
        
    except Exception as e:
        return None, str(e)
